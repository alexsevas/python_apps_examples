from pathlib import Path
from pypdf import PdfReader
from docx import Document


def load_text(path):
    """Загружает текст из файла с обработкой ошибок и ограничением размера"""
    path = Path(path)
    
    # Проверка размера файла (максимум 50MB)
    max_size = 50 * 1024 * 1024  # 50MB
    if path.stat().st_size > max_size:
        raise Exception(f"File too large. Maximum size is 50MB, file is {path.stat().st_size / 1024 / 1024:.1f}MB")
    
    if path.suffix.lower() == ".txt":
        # Пробуем разные кодировки
        for encoding in ['utf-8', 'utf-8-sig', 'cp1251', 'latin-1']:
            try:
                text = path.read_text(encoding=encoding)
                if text.strip():  # Проверяем что текст не пустой
                    return text
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise Exception("Failed to decode text file. Unsupported encoding.")

    if path.suffix.lower() == ".pdf":
        try:
            reader = PdfReader(path)
            if not reader.pages:
                raise Exception("PDF file has no pages")
            
            text = "\n".join(
                page.extract_text() or ""
                for page in reader.pages
            )
            
            if not text.strip():
                raise Exception("PDF file contains no extractable text")
            
            return text
        except Exception as e:
            raise Exception(f"Failed to read PDF: {str(e)}")

    if path.suffix.lower() == ".docx":
        try:
            doc = Document(path)
            if not doc.paragraphs:
                raise Exception("DOCX file has no paragraphs")
            
            text = "\n".join(
                p.text for p in doc.paragraphs
            )
            
            if not text.strip():
                raise Exception("DOCX file contains no text")
            
            return text
        except Exception as e:
            raise Exception(f"Failed to read DOCX: {str(e)}")

    raise Exception(f"Unsupported file format: {path.suffix}")
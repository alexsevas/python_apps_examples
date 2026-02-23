# pip install python-docx

from docx import Document
from docx.shared import Inches

guestnames=['Иван', 'Олег', 'Ольга', 'Мария и Александр']

for x in guestnames:
    document = Document()
    document.add_heading('Приглашение на свадьбу!', 0)
    p = document.add_paragraph(x+', спешим пригласить Вас на свадьбу')
    p.add_run(' Виктора и Инны Факсмейт').bold = True
    p.add_run(', которая состоится 12.05.2014 в Питере.')
    document.add_heading('Мы ждём вас, на наш праздник!!!', level=1)
    document.add_paragraph('')
    document.add_picture('1.jpg', width=Inches(5.25))
    document.add_page_break()
    document.save(x+'.docx')
